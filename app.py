import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO
import sqlite3

# 1. CONFIGURATION DE LA PAGE
st.set_page_config(page_title="Instructions Compositeur", layout="centered")

# 2. GESTION DU SYSTÈME BILINGUE
if "langue" not in st.session_state:
    st.session_state.langue = "Français"

# Sélecteur de langue placé discrètement en haut de l'écran
col_titre, col_langue = st.columns([3, 1])
with col_langue:
    st.session_state.langue = st.selectbox(
        "🌐 Language :", 
        ["Français", "English"], 
        index=0 if st.session_state.langue == "Français" else 1
    )

# Dictionnaire centralisé de toutes les traductions de l'interface
TRAD = {
    "Français": {
        "titre_app": "Instructions de mise en page",
        "tab_editeurs": "✍️ Éditeurs",
        "tab_compositeurs": "🎼 Compositeurs",
        "titre_editeur": "Espace Éditeurs",
        "titre_compositeur": "Espace Compositeurs",
        "auth_label": "Veuillez saisir le mot de passe pour accéder à cet espace :",
        "auth_btn": "Se connecter",
        "auth_err": "🔑 Mot de passe incorrect.",
        "logout_btn": "🔒 Se déconnecter de l'espace Éditeur",
        "choix_revue": "Choisir ou chercher une revue :",
        "choix_defaut": "-- Sélectionnez une revue --",
        "btn_word": "📄 Télécharger au format Word (.docx)",
        "msg_attente": "ℹ️ L'application est prête. Veuillez vous connecter à l'espace Éditeur pour y charger votre fichier Excel d'origine.",
        "save_success": "Base locale mise à jour !",
        "save_btn": "💾 Enregistrer et appliquer définitivement"
    },
    "English": {
        "titre_app": "Layout Instructions",
        "tab_editeurs": "✍️ Editors",
        "tab_compositeurs": "🎼 Compositors",
        "titre_editeur": "Editors Workspace",
        "titre_compositeur": "Compositors Workspace",
        "auth_label": "Please enter the password to access this area:",
        "auth_btn": "Login",
        "auth_err": "🔑 Incorrect password.",
        "logout_btn": "🔒 Log out from Editors Workspace",
        "choix_revue": "Select or search for a journal:",
        "choix_defaut": "-- Select a journal --",
        "btn_word": "📄 Download in Word format (.docx)",
        "msg_attente": "ℹ️ The app is ready. Please log into the Editor workspace to upload your original Excel file.",
        "save_success": "Local database updated!",
        "save_btn": "💾 Save changes permanently"
    }
}

# Raccourci pour utiliser les textes dans la langue active
T = TRAD[st.session_state.langue]

st.title(T["titre_app"])

MOT_DE_PASSE_EDITEUR = "Editeur2026"
DB_NOM = "revues.db"

# [Le reste de vos fonctions d'initialisation SQLite reste identique...]
def initialiser_sqlite():
    conn = sqlite3.connect(DB_NOM)
    cursor = conn.cursor()
    cursor.execute("CREATE TABLE IF NOT EXISTS instructions (revue TEXT PRIMARY KEY, donnees_json TEXT);")
    conn.commit()
    conn.close()

initialiser_sqlite()

def charger_donnees_sqlite():
    conn = sqlite3.connect(DB_NOM)
    try:
        df_sql = pd.read_sql_query("SELECT * FROM instructions", conn)
        conn.close()
        if df_sql.empty:
            return None
        import json, liste_dictionnaires
        liste_dictionnaires = []
        for _, row in df_sql.iterrows():
            dict_revue = json.loads(row["donnees_json"])
            dict_revue["Revue"] = row["revue"]
            liste_dictionnaires.append(dict_revue)
        df_final = pd.DataFrame(liste_dictionnaires)
        return df_final.set_index("Revue")
    except Exception:
        conn.close()
        return None

def sauvegarder_revue_sqlite(nom_revue, dictionnaire_sections):
    import json
    conn = sqlite3.connect(DB_NOM)
    cursor = conn.cursor()
    json_consignes = json.dumps(dictionnaire_sections)
    cursor.execute("""
        INSERT INTO instructions (revue, donnees_json) VALUES (?, ?)
        ON CONFLICT(revue) DO UPDATE SET donnees_json = excluded.donnees_json;
    """, (nom_revue, json_consignes))
    conn.commit()
    conn.close()

def supprimer_revue_sqlite(nom_revue):
    conn = sqlite3.connect(DB_NOM)
    cursor = conn.cursor()
    cursor.execute("DELETE FROM instructions WHERE revue = ?;", (nom_revue,))
    conn.commit()
    conn.close()

def generer_document_word(nom_revue, données_instructions):
    doc = Document()
    doc.add_heading(f"Instructions de mise en page — {nom_revue}" if st.session_state.langue == "Français" else f"Layout Instructions — {nom_revue}", level=1)
    for section, contenu in données_instructions.items():
        if pd.notna(contenu) and str(contenu).strip() not in ["", "/"]:
            doc.add_heading(str(section).capitalize(), level=2)
            doc.add_paragraph(str(contenu).strip())
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

st.session_state.df_revues = charger_donnees_sqlite()

# Remplacement des onglets par les variables traduites
tab_editeurs, tab_compositeurs = st.tabs([T["tab_editeurs"], T["tab_compositeurs"]])

# ==========================================
# 1. POINT D'ENTRÉE : ÉDITEURS (SUITE DU BLOC 1)
# ==========================================
with tab_editeurs:
    st.header(T["titre_editeur"])
    
    if not st.session_state.authentifie:
        with st.form("form_auth"):
            mdp_saisi = st.text_input(T["auth_label"], type="password")
            valider_auth = st.form_submit_button(T["auth_btn"])
            if valider_auth:
                if mdp_saisi == MOT_DE_PASSE_EDITEUR:
                    st.session_state.authentifie = True
                    st.rerun()
                else:
                    st.error(T["auth_err"])
    else:
        # Bouton de déconnexion bilingue
        if st.button(T["logout_btn"]):
            st.session_state.authentifie = False
            st.rerun()
            
        st.markdown("---")

        if st.session_state.df_revues is not None:
            df_edition = st.session_state.df_revues
            liste_sections = list(df_edition.columns)
            
            # --- SECTION 1 : MODIFIER ET SAUVEGARDER EN DIRECT ---
            lbl_modif_titre = "1. Modifier et Sauvegarder en direct" if st.session_state.langue == "Français" else "1. Edit and Save Live"
            lbl_radio = ["Pour une seule revue", "Pour l'ensemble des revues"] if st.session_state.langue == "Français" else ["For a single journal", "For all journals"]
            
            st.subheader(lbl_modif_titre)
            mode_modification = st.radio(
                "Périmètre :" if st.session_state.langue == "Français" else "Scope :",
                lbl_radio,
                horizontal=True
            )
            
            # CAS 1 : MODIFICATION UNIQUE
            if mode_modification in ["Pour une seule revue", "For a single journal"]:
                lbl_select_rev = "Sélectionner la revue à éditer :" if st.session_state.langue == "Français" else "Select the journal to edit:"
                revue_a_modifier = st.selectbox(lbl_select_rev, df_edition.index)
                
                with st.form("form_mono_revue"):
                    st.write(f"✍️ Modifications : **{revue_a_modifier}**")
                    nouveaux_contenus = {}
                    
                    for section in liste_sections:
                        valeur_actuelle = str(df_edition.loc[revue_a_modifier, section])
                        if valeur_actuelle == "nan" or valeur_actuelle == "/":
                            valeur_actuelle = ""
                        
                        nouveaux_contenus[section] = st.text_area(f"Section : {section}", value=valeur_actuelle)
                    
                    soumettre = st.form_submit_button(T["save_btn"])
                    if soumettre:
                        dict_sauvegarde = {}
                        for section, nv_texte in nouveaux_contenus.items():
                            dict_sauvegarde[section] = nv_texte if nv_texte.strip() != "" else "/"
                        
                        sauvegarder_revue_sqlite(revue_a_modifier, dict_sauvegarde)
                        st.toast(T["save_success"], icon="💾")
                        st.rerun()

            # CAS 2 : MODIFICATION GLOBALE
            else:
                lbl_select_sec = "Sélectionner la section à harmoniser partout :" if st.session_state.langue == "Français" else "Select the section to standardize everywhere:"
                section_a_modifier = st.selectbox(lbl_select_sec, liste_sections)
                premiere_revue_nom = df_edition.index if len(df_edition.index) > 0 else "Aucune"
                
                valeur_premiere_revue = ""
                if len(df_edition.index) > 0:
                    valeur_premiere_revue = str(df_edition.at[premiere_revue_nom, section_a_modifier])
                
                if valeur_premiere_revue == "nan" or valeur_premiere_revue == "/":
                    valeur_premiere_revue = ""
                
                with st.form("form_global_revue"):
                    txt_alerte = f"🚨 Vous allez écraser la section **{section_a_modifier}** pour **toutes** les revues." if st.session_state.langue == "Français" else f"🚨 You are about to overwrite the **{section_a_modifier}** section for **all** journals."
                    txt_caption = f"💡 Champ pré-rempli avec le texte actuel de : *{premiere_revue_nom}*." if st.session_state.langue == "Français" else f"💡 Field pre-filled with current text from: *{premiere_revue_nom}*."
                    lbl_text_area = "Nouveau texte commun :" if st.session_state.langue == "Français" else "New common text:"
                    btn_global = "⚠️ Écraser et Sauvegarder sur tout le catalogue" if st.session_state.langue == "Français" else "⚠️ Overwrite and Save across catalog"
                    
                    st.write(txt_alerte)
                    st.caption(txt_caption)
                    
                    texte_global = st.text_area(lbl_text_area, value=valeur_premiere_revue)
                    soumettre_global = st.form_submit_button(btn_global)
                    
                    if soumettre_global:
                        texte_global_propre = texte_global if texte_global.strip() != "" else "/"
                        
                        for nom_revue in df_edition.index:
                            dict_actuel = df_edition.loc[nom_revue].to_dict()
                            dict_actuel[section_a_modifier] = texte_global_propre
                            sauvegarder_revue_sqlite(nom_revue, dict_actuel)
                            
                        st.toast(T["save_success"], icon="💾")
                        st.rerun()
            # --- SECTION 2 : STRUCTURER LA BASE DE DONNÉES (AJOUT & SUPPRESSION) ---
            lbl_struct_titre = "2. Structurer la base de données" if st.session_state.langue == "Français" else "2. Database Structure"
            st.markdown("---")
            st.subheader(lbl_struct_titre)
            
            # Bloc A : AJOUTS
            st.markdown("##### ➕ " + ("Ajouts" if st.session_state.langue == "Français" else "Additions"))
            col_revue, col_section = st.columns(2)
            
            with col_revue:
                with st.form("form_ajouter_revue"):
                    st.write("**" + ("Ajouter une nouvelle revue" if st.session_state.langue == "Français" else "Add a new journal") + "**")
                    lbl_in_rev = "Nom de la nouvelle revue :" if st.session_state.langue == "Français" else "New journal name:"
                    btn_in_rev = "Créer la revue" if st.session_state.langue == "Français" else "Create journal"
                    nouvelle_revue_nom = st.text_input(lbl_in_rev)
                    soumettre_nouvelle_revue = st.form_submit_button(btn_in_rev)
                    if soumettre_nouvelle_revue and nouvelle_revue_nom.strip() != "":
                        nom_propre = nouvelle_revue_nom.strip()
                        if nom_propre in st.session_state.df_revues.index:
                            st.error("⚠️ Existe déjà." if st.session_state.langue == "Français" else "⚠️ Already exists.")
                        else:
                            dict_vide = {col: "/" for col in liste_sections}
                            sauvegarder_revue_sqlite(nom_propre, dict_vide)
                            st.rerun()
                                
            with col_section:
                with st.form("form_ajouter_section"):
                    st.write("**" + ("Ajouter une nouvelle section" if st.session_state.langue == "Français" else "Add a new section") + "**")
                    lbl_in_sec = "Nom de la nouvelle section :" if st.session_state.langue == "Français" else "New section name:"
                    btn_in_sec = "Créer la section" if st.session_state.langue == "Français" else "Create section"
                    nouvelle_section_nom = st.text_input(lbl_in_sec)
                    soumettre_nouvelle_section = st.form_submit_button(btn_in_sec)
                    if soumettre_nouvelle_section and nouvelle_section_nom.strip() != "":
                        sec_propre = nouvelle_section_nom.strip()
                        if sec_propre in st.session_state.df_revues.columns:
                            st.error("⚠️ Existe déjà." if st.session_state.langue == "Français" else "⚠️ Already exists.")
                        else:
                            for nom_revue in df_edition.index:
                                dict_actuel = df_edition.loc[nom_revue].to_dict()
                                dict_actuel[sec_propre] = "/"
                                sauvegarder_revue_sqlite(nom_revue, dict_actuel)
                            st.rerun()

            # Bloc B : SUPPRESSIONS
            st.markdown("##### 🗑️ " + ("Suppressions (Irréversible)" if st.session_state.langue == "Français" else "Deletions (Irreversible)"))
            col_del_revue, col_del_section = st.columns(2)
            
            with col_del_revue:
                with st.form("form_supprimer_revue"):
                    st.write("**" + ("Supprimer une revue" if st.session_state.langue == "Français" else "Delete a journal") + "**")
                    lbl_drop_rev = "Revue à détruire :" if st.session_state.langue == "Français" else "Journal to destroy:"
                    btn_drop_rev = "💥 Supprimer" if st.session_state.langue == "Français" else "💥 Delete"
                    revue_a_supprimer = st.selectbox(lbl_drop_rev, ["-- Sélectionner/Select --"] + list(df_edition.index))
                    soumettre_del_revue = st.form_submit_button(btn_drop_rev)
                    if soumettre_del_revue and revue_a_supprimer != "-- Sélectionner/Select --":
                        supprimer_revue_sqlite(revue_a_supprimer)
                        st.rerun()
                            
            with col_del_section:
                with st.form("form_supprimer_section"):
                    st.write("**" + ("Supprimer une section complète" if st.session_state.langue == "Français" else "Delete a full section") + "**")
                    lbl_drop_sec = "Section à détruire :" if st.session_state.langue == "Français" else "Section to destroy:"
                    btn_drop_sec = "💥 Supprimer" if st.session_state.langue == "Français" else "💥 Delete"
                    section_a_supprimer = st.selectbox(lbl_drop_sec, ["-- Sélectionner/Select --"] + liste_sections)
                    soumettre_del_section = st.form_submit_button(btn_drop_sec)
                    if soumettre_del_section and section_a_supprimer != "-- Sélectionner/Select --":
                        for nom_revue in df_edition.index:
                            dict_actuel = df_edition.loc[nom_revue].to_dict()
                            if section_a_supprimer in dict_actuel:
                                del dict_actuel[section_a_supprimer]
                            sauvegarder_revue_sqlite(nom_revue, dict_actuel)
                        st.rerun()

        # --- SECTION 3 : IMPORT INITIAL (ANTI-BOUCLE) ---
        lbl_imp_titre = "3. Remplissage ou Remplacement global via Excel" if st.session_state.langue == "Français" else "3. Global Upload or Replacement via Excel"
        lbl_imp_desc = "Déposez votre fichier Excel pour pré-remplir la base de données." if st.session_state.langue == "Français" else "Upload your Excel file to pre-fill the database."
        lbl_file_pld = "Déposer le fichier Excel (.xlsx)" if st.session_state.langue == "Français" else "Drop the Excel file (.xlsx)"
        
        st.markdown("---")
        st.subheader(lbl_imp_titre)
        st.write(lbl_imp_desc)
        
        if "import_deja_fait" not in st.session_state:
            st.session_state.import_deja_fait = False

        fichier_charge = st.file_uploader(lbl_file_pld, type=["xlsx"], key="uploader_excel")
        if fichier_charge is None:
            st.session_state.import_deja_fait = False
        
        if fichier_charge is not None and not st.session_state.import_deja_fait:
            try:
                df_nouveau = pd.read_excel(fichier_charge)
                if "Revue" in df_nouveau.columns:
                    conn_clear = sqlite3.connect(DB_NOM)
                    conn_clear.execute("DELETE FROM instructions;")
                    conn_clear.commit()
                    conn_clear.close()
                    
                    for _, row in df_nouveau.iterrows():
                        nom_revue = str(row["Revue"]).strip()
                        dict_revue = {}
                        for col in df_nouveau.columns:
                            if col != "Revue":
                                dict_revue[str(col).strip()] = str(row[col]).strip() if pd.notna(row[col]) else "/"
                        sauvegarder_revue_sqlite(nom_revue, dict_revue)
                    
                    st.session_state.import_deja_fait = True
                    st.success("✅ Succès !" if st.session_state.langue == "Français" else "✅ Success!")
                    st.rerun()
                else:
                    st.error("⚠️ Colonne 'Revue' manquante." if st.session_state.langue == "Français" else "⚠️ Missing 'Revue' column.")
            except Exception as e:
                st.error(f"⚠️ Erreur : {e}")

# ==========================================
# 2. POINT D'ENTRÉE : COMPOSITEURS
# ==========================================
with tab_compositeurs:
    st.header(T["titre_compositeur"])
    
    if st.session_state.df_revues is None:
        st.info(T["msg_attente"])
    else:
        df_revues = st.session_state.df_revues
        
        choix = st.selectbox(
            T["choix_revue"], 
            [T["choix_defaut"]] + list(df_revues.index),
            key="select_compositeur"
        )

        if choix != T["choix_defaut"]:
            instructions_revue = df_revues.loc[choix]
            
            fichier_word = generer_document_word(choix, instructions_revue)
            st.download_button(
                label=T["btn_word"],
                data=fichier_word,
                file_name=f"Instructions_{choix}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
            st.markdown("---")
            
            for section, contenu in instructions_revue.items():
                if pd.notna(contenu) and str(contenu).strip() not in ["", "/"]:
                    st.subheader(str(section))
                    st.write(str(contenu).strip())
